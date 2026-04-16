import os
import uuid
import concurrent.futures

# Diagnostic Import Checks
IMPORT_ERRORS = []
try:
    import fitz
except Exception as e:
    IMPORT_ERRORS.append(f"PyMuPDF (fitz) error: {str(e)}")

try:
    import pytesseract
except Exception as e:
    IMPORT_ERRORS.append(f"Pytesseract error: {str(e)}")

try:
    from docx import Document
except Exception as e:
    IMPORT_ERRORS.append(f"python-docx error: {str(e)}")

from flask import Flask, render_template, request, jsonify, send_file
from PIL import Image
import traceback
import requests
import base64
import io
import time

# Absolute Paths for stability
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Environment Detection
IS_VERCEL = "VERCEL" in os.environ or os.environ.get("VERCEL_REGION") is not None

# Tesseract Path Configuration
if os.name == 'nt':  # Windows Local
    TESS_PATH = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    pytesseract.pytesseract.tesseract_cmd = TESS_PATH
else:
    # On Linux/Vercel, we assume tesseract is in the PATH
    pass

# Direct relative path for templates
app = Flask(__name__, template_folder='templates')

# Progress Tracking
processing_status = {}

@app.route('/status/<file_id>')
def get_status(file_id):
    return jsonify(processing_status.get(file_id, {"current": 0, "total": 0, "status": "processing"}))

if IS_VERCEL:
    UPLOAD_FOLDER = '/tmp/uploads'
    PROCESSED_FOLDER = '/tmp/processed'
else:
    UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
    PROCESSED_FOLDER = os.path.join(os.path.expanduser("~"), "Downloads")

def ensure_dirs():
    """Ensure upload and processed directories exist before use."""
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(PROCESSED_FOLDER, exist_ok=True)

def process_single_page(args):
    input_pdf_path, page_num, file_id = args
    temp_pdf_page = os.path.join(PROCESSED_FOLDER, f"{file_id}_p{page_num}.pdf")
    
    try:
        if not os.path.exists(input_pdf_path):
            return None, "File missing"
            
        print(f"[*] Starting OCR: Page {page_num+1}...")
        # Open the specific page inside the worker to save overall memory
        doc = fitz.open(input_pdf_path)
        page = doc[page_num]
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        doc.close() # Close immediately after rendering

        # Local Tesseract Attempt
        try:
            # Generate searchable PDF for this page
            pdf_data = pytesseract.image_to_pdf_or_hocr(img, extension='pdf')
            with open(temp_pdf_page, 'wb') as f:
                f.write(pdf_data)
            
            # Extract text for Word doc
            text = pytesseract.image_to_string(img)
            print(f"[OK] Page {page_num+1} (Local) Completed.")
            return temp_pdf_page, text
        except Exception as local_err:
            # Cloud Fallback (OCR.space)
            print(f"[*] Local OCR Failed, trying Cloud for Page {page_num+1}...")
            # Convert image to bytes for API
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format='JPEG')
            img_bytes = img_byte_arr.getvalue()
            
            payload = {
                'apikey': 'helloworld', # Default free key
                'language': 'eng',
                'isOverlayRequired': False,
                'isCreateSearchablePdf': False, # Free tier doesn't support searchable PDF well for single images
            }
            
            # Send to OCR.space
            files = {'file': ('page.jpg', img_bytes, 'image/jpeg')}
            response = requests.post('https://api.ocr.space/parse/image', files=files, data=payload, timeout=20)
            
            if response.status_code != 200:
                raise Exception(f"API Server Error: {response.status_code}")
                
            result = response.json()
            
            # Check if result is a dictionary (it should be)
            if isinstance(result, dict) and result.get('OCRExitCode') == 1:
                parsed_results = result.get('ParsedResults', [])
                if parsed_results:
                    text = parsed_results[0].get('ParsedText', '')
                    print(f"[OK] Page {page_num+1} (Cloud) Completed.")
                    return None, text
                else:
                    raise Exception("API returned success but no text found.")
            else:
                error_info = result.get('ErrorMessage') if isinstance(result, dict) else str(result)
                raise Exception(f"Cloud API Error: {error_info}")

    except Exception as e:
        error_msg = str(e)
        print(f"[!] Error on Page {page_num+1}: {error_msg}")
        return None, f"Page {error_msg}"

def master_ocr_process(input_pdf_path, file_id):
    # Only open to get page count, then close
    doc_meta = fitz.open(input_pdf_path)
    total_pages = len(doc_meta)
    doc_meta.close()
    
    print(f"\n--- Processing Document: {total_pages} Pages total ---")

    worker_args = [(input_pdf_path, i, file_id) for i in range(total_pages)]

    temp_pdfs = []
    full_text = []

    print(f"Running OCR Engine (Memory Optimized Mode)...")
    processing_status[file_id] = {"current": 0, "total": total_pages, "status": "processing"}
    
    # On Vercel (Cloud Mode), we MUST use 1 worker because free API keys only allow 1 concurrent request
    max_workers = 1 if IS_VERCEL else 3
    
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(process_single_page, arg): i for i, arg in enumerate(worker_args)}
        
        # We need to preserve order for the final merge/docx, but results come in any order
        ordered_results = [None] * total_pages
        completed_count = 0
        
        for future in concurrent.futures.as_completed(futures):
            original_index = futures[future]
            try:
                res = future.result()
                ordered_results[original_index] = res
                completed_count += 1
                processing_status[file_id] = {"current": completed_count, "total": total_pages, "status": "processing"}
            except Exception as e:
                print(f"Future error: {e}")
                ordered_results[original_index] = (None, "")
                completed_count += 1

    last_worker_error = "Unknown"
    # Extract results into ordered lists
    for res in ordered_results:
        if res and res[0]:
            temp_pdfs.append(res[0])
            full_text.append(res[1])
        elif res and res[1]:
            last_worker_error = res[1]

    if not temp_pdfs and not full_text:
        raise Exception(f"OCR failed to process any pages. Last error from engine: {last_worker_error}")

    # Initialize output paths
    output_pdf_path = None
    output_docx_path = None

    if temp_pdfs:
        print(f"Merging {len(temp_pdfs)} pages into final Searchable PDF...")
        output_pdf_path = os.path.join(PROCESSED_FOLDER, f"{file_id}_searchable.pdf")
        combined_pdf = fitz.open()
        
        for temp_pdf in temp_pdfs:
            if os.path.exists(temp_pdf):
                with fitz.open(temp_pdf) as m_pdf:
                    combined_pdf.insert_pdf(m_pdf)
                os.remove(temp_pdf)

        # Save with high compression and garbage collection to reduce file size
        combined_pdf.save(output_pdf_path, garbage=3, deflate=True, clean=True)
        combined_pdf.close()
    else:
        print("[!] No searchable PDF pages generated (Cloud fallback active).")

    if full_text:
        print("Creating Word Document...")
        output_docx_path = os.path.join(PROCESSED_FOLDER, f"{file_id}_editable.docx")
        doc = Document()
        for i, text in enumerate(full_text):
            doc.add_heading(f'Page {i+1}', level=1)
            doc.add_paragraph(text)
            if i < len(full_text) - 1:
                doc.add_page_break()
        doc.save(output_docx_path)
    
    print("--- Document Processing Completed ---\n")
    processing_status[file_id]["status"] = "completed"
    return output_pdf_path, output_docx_path

@app.route('/')
def index():
    status_msg = ""
    if IMPORT_ERRORS:
        status_msg = f"<div style='background:#fee2e2; border:1px solid #ef4444; color:#991b1b; padding:10px; margin-bottom:20px; border-radius:10px; font-size:12px;'><strong>Vercel Compatibility Note:</strong> Local libraries (Tesseract/Poppler) are missing. System will use Cloud OCR Fallback.</div>"
    
    try:
        return render_template('index.html', status_msg=status_msg)
    except Exception as e:
        return f"<h1>Template Load Error</h1><p>{str(e)}</p>", 500

@app.errorhandler(500)
def handle_500(e):
    return f"<h1>Internal Server Error (Vercel Debug)</h1><pre>{traceback.format_exc()}</pre>", 500

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        ensure_dirs() # Create folders only when needed
        file = request.files['file']
        file_id = request.form.get('file_id', str(uuid.uuid4()))
        original_path = os.path.join(UPLOAD_FOLDER, f"{file_id}_orig.pdf")
        file.save(original_path)

        # Start master process
        pdf, docx = master_ocr_process(original_path, file_id)
        
        if os.path.exists(original_path): os.remove(original_path)
        return jsonify({'success': True, 'file_id': file_id})
    except Exception as e:
        print(f"Global Error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/download/<file_id>/<filename>')
def download_file(file_id, filename):
    folder = PROCESSED_FOLDER
    # Determine which file to serve based on the filename requested
    if filename.endswith('.pdf'):
        internal_name = f"{file_id}_searchable.pdf"
    elif filename.endswith('.docx'):
        internal_name = f"{file_id}_editable.docx"
    else:
        return "Invalid file type", 400
        
    path = os.path.join(folder, internal_name)
    if os.path.exists(path):
        return send_file(path, as_attachment=True, download_name=filename)
    return "File not found on server", 404

if __name__ == '__main__':
    # Local development - Keep debug False to prevent restarts during long OCR tasks
    app.run(debug=False, port=int(os.environ.get("PORT", 5000)))
else:
    # Production (Vercel)
    app.debug = False
